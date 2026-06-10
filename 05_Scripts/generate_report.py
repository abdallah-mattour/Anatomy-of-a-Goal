# -*- coding: utf-8 -*-
"""Generate the formal 3-5 page Word report for the AI-Assisted Data Analysis assignment."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Project root = the folder that holds 01_Report, 02_Notebook, ... (parent of 05_Scripts)
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHARTS = os.path.join(ROOT, '03_Figures')
OUT = os.path.join(ROOT, '01_Report', 'Manhaji_Report_Anatomy_of_a_Goal.docx')

# Brand colors
GREEN = RGBColor(0x2D, 0x57, 0x2C)
DARKBLUE = RGBColor(0x1F, 0x3A, 0x5F)
RED = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- Base style ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.05

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading(text, size=15, color=GREEN, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Calibri'
    return p

def add_body(text, size=10.5, italic=False, color=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p

def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p

def add_image(fname, width=6.2, caption=None):
    path = os.path.join(CHARTS, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_after = Pt(8)
            cr = c.add_run(caption)
            cr.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = GRAY

# ============ TITLE BLOCK ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
tr = title.add_run('The Anatomy of a Goal')
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = GREEN
tr.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
sr = sub.add_run('Decoding Scoring Patterns Across Europe’s Top Leagues')
sr.font.size = Pt(13)
sr.font.color.rgb = DARKBLUE
sr.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_after = Pt(8)
sr2 = sub2.add_run('AI-Assisted Data Analysis Using Agentic AI  —  Report')
sr2.font.size = Pt(11)
sr2.bold = True

# Divider line
divp = doc.add_paragraph()
divp.paragraph_format.space_after = Pt(8)
pPr = divp._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '2D572C')
pbdr.append(bottom)
pPr.append(pbdr)

# Info line
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_after = Pt(10)
ir = info.add_run('Dataset: Wyscout Soccer Match Event Dataset (Kaggle)   |   AI Tool: Claude (Anthropic)   |   Analysis: Python, pandas, matplotlib, seaborn')
ir.font.size = Pt(9)
ir.font.color.rgb = GRAY

# ============ TASK 1: DATASET DESCRIPTION ============
add_heading('1. Dataset Description', size=15)

add_body(
    'This study analyzes the Wyscout Soccer Match Event Dataset, a professional-grade football '
    'analytics resource published on Kaggle (aleespinosa/soccer-match-event-dataset). The data captures '
    'every on-ball action from 1,941 matches played during the 2017/2018 club season and the 2016 European '
    'Championship and 2018 World Cup. The raw events were converted into the SPADL format (Soccer Player '
    'Action Description Language), which standardizes each action with pitch coordinates, an action type, '
    'a result, and the body part used.'
)

add_body('The analysis uses three row-aligned core files plus six supporting reference tables:', bold=False)

# Source/scale table
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for c, txt in zip(hdr, ['File', 'Rows × Columns', 'Description']):
    c.paragraphs[0].add_run(txt).bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9.5)
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_bg(c, '2D572C')

rows_data = [
    ('actions.csv', '2,462,726 × 17', 'Every on-ball action: x/y coordinates, type, result, body part'),
    ('features.csv', '2,462,726 × 86', 'Engineered spatial features for 3-action sequences (distance, angle)'),
    ('labels.csv', '2,462,726 × 2', 'VAEP labels: "scores" and "concedes" (does the sequence lead to a goal?)'),
    ('games.csv', '1,941 × 8', 'Match reference: competition, date, home/away teams'),
    ('players / teams', '3,603 / 142', 'Player profiles (position, foot, height) and club metadata'),
    ('playerank.csv', '46,897 × 6', 'Per-match player ratings with positional role clusters'),
]
for fname, dims, desc in rows_data:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(fname).font.size = Pt(9)
    cells[0].paragraphs[0].runs[0].bold = True
    cells[1].paragraphs[0].add_run(dims).font.size = Pt(9)
    cells[2].paragraphs[0].add_run(desc).font.size = Pt(9)

# widths
for row in tbl.rows:
    row.cells[0].width = Inches(1.4)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(3.6)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_body(
    'Main variables. The most analytically valuable columns are the spatial coordinates '
    '(start_x, start_y, end_x, end_y on a 105 × 68 m pitch), the action type_name (19 distinct '
    'types including pass, cross, dribble, shot, interception, and tackle), the result_name '
    '(success / fail / offside / owngoal), the bodypart_name (foot / head / other), and the engineered '
    'features start_distance_to_goal and start_angle_to_goal. The binary scores label flags the actions '
    'that belong to a goal-scoring sequence. In total the dataset contains 2,462,726 actions, of which 45,937 are shots and '
    '5,104 are goals — an overall conversion rate of 11.1% and an average of 2.63 goals per match.'
)

# ============ TASK 2: AI-ASSISTED ANALYSIS ============
add_heading('2. AI-Assisted Analysis', size=15)

add_body(
    'The full analysis was conducted in a Jupyter Notebook generated and iteratively refined with the help '
    'of Claude (Anthropic). Claude assisted in four phases: (a) exploring and joining the 27 raw files, '
    '(b) detecting and explaining missing values, (c) computing descriptive statistics and 13 visualizations, '
    'and (d) interpreting the spatial, sequential and cross-league patterns behind goals. The example prompts used included '
    '“Analyze this dataset and summarize the main insights,” “Detect missing values and recommend '
    'preprocessing steps,” and “Generate charts and identify the most important variables.”'
)

add_heading('2.1 Missing Values & Preprocessing', size=12, color=DARKBLUE, space_before=6)
add_body(
    'Missing-value detection showed the dataset is remarkably complete. Only original_event_id had a '
    'notable gap (6.85%), and the sequence-context columns (suffixed -1 and -2) were missing in just '
    '0.08–0.16% of rows. Crucially, these gaps are not data corruption: they correspond to the first '
    'and second actions of each match half, which have no preceding action to describe. The recommended '
    'preprocessing — filling these context columns with 0 and treating the success/fail result as the '
    'outcome of interest — was applied before the analysis.'
)

add_heading('2.2 Where Goals Come From', size=12, color=DARKBLUE, space_before=6)
add_body(
    'Mapping all 5,104 goals onto the pitch reveals an unmistakable concentration: goals are scored from a '
    'narrow central corridor directly in front of the goal, overwhelmingly inside the penalty area. The '
    'build-up heatmap (Chart 2) shows scoring sequences gather most density just outside the box before the '
    'decisive entry.'
)
add_image('01_shot_location_heatmap.png', width=4.3,
          caption='Chart 1 — Goal density: nearly all goals originate within ~18 m of the goal, centrally.')

add_heading('2.3 The Winning Formula: Sequences That Score', size=12, color=DARKBLUE, space_before=6)
add_body(
    'A goal is the climax of a sequence, not an isolated event. Examining the three consecutive actions that '
    'end in a goal, the single most common chain is pass → pass → shot (831 goals), followed by '
    'pass → cross → shot (542) and pass → dribble → shot (436). Three signatures stand out: a passing core, '
    'a wide-delivery route in which a cross feeds the shot, and a dribble-led route where a player beats an '
    'opponent before shooting or assisting. Two further patterns are revealing — shot → save → shot (225 goals) '
    'shows the value of following up rebounds, while sequences that win a foul in the box convert into '
    'penalties. Across all chains, shots are 9.1×, crosses 2.45×, and dribbles 1.49× over-represented in '
    'goal-scoring sequences versus normal play.'
)
add_image('04_top_action_sequences.png', width=5.8,
          caption='Chart 4 — The 15 most common three-action sequences that end in goals, colour-coded by play type.')

add_heading('2.4 League Fingerprints: How Scoring Differs Across Europe', size=12, color=DARKBLUE, space_before=6)
add_body(
    'Each competition has a distinct scoring signature. The 2018 World Cup was the most prolific at 2.86 '
    'goals per match, with the Bundesliga the highest-scoring domestic league (2.72). The starkest contrast '
    'is the source of goals: in the international tournaments, penalties account for 26–27% of goals (Euro '
    '27.1%, World Cup 26.2%) — three to four times their 6–10% share in the domestic leagues, an effect '
    'inflated by penalty shootouts in the knockout rounds. Open-play creation differs too: Ligue 1 (19.3%) '
    'and La Liga (19.0%) lean most heavily on crosses, whereas the World Cup is the least cross-dependent '
    '(11.5%), favouring more varied, direct attacks. Headed goals stay broadly consistent (14–18%), while '
    'average shot distance is slightly longer in the international tournaments.'
)
add_image('13_league_comparison.png', width=6.3,
          caption='Chart 13 — Scoring DNA: goals per match, cross-assisted %, penalty %, and headed % across the seven competitions.')

# ============ TASK 3: INSIGHTS ============
add_heading('3. Key Insights', size=15)
add_body('Six data-backed findings emerged. Each was computed directly from the data, stress-tested for robustness, and cross-checked against football intuition.', italic=True, size=10)

insights = [
    ('Insight 1 — The Golden Zone. ',
     '84.4% of all goals come from within 18 m of goal, where shots convert at 18.1% versus just 3.6% from '
     'outside — a 5× edge. Closer still, the point-blank zone (≤11 m) alone supplies 46% of all goals at a '
     '28% conversion rate. The lesson is blunt: work the ball into the box rather than shoot from distance.'),
    ('Insight 2 — The Cross-Shot Pipeline. ',
     'A cross is the most efficient final ball in football: crosses convert at 16.6% versus 9.3% for a normal '
     'pass — nearly twice as likely to produce a goal — and they supply 18% of all goals. Delivery skews to '
     'one flank, with the right wing creating more than the left (503 vs 405).'),
    ('Insight 3 — The Cut-Back Is King. ',
     'Roughly 19% of goals are created by a cut-back: a player carries the ball to the deep, wide byline and '
     'pulls it back to a teammate arriving centrally. It is one of the most repeatable, coachable patterns in '
     'the modern game, and it stands out clearly in the assist-delivery map (Chart 3).'),
    ('Insight 4 — Two Routes Into the Box. ',
     'The decisive pass is almost always played close: 64.6% of assists are delivered from inside the penalty '
     'area, a median of just 20.3 m from goal. That supply splits into two complementary routes — wide '
     'deliveries (57% of cross-assists come from the flanks) and central short passes (43% played from inside '
     'the box).'),
    ('Insight 5 — Formation Vulnerability. ',
     'Among formations used in at least 100 matches, goals conceded range from 1.01 per match (a 2-3-5 shape) '
     'to 1.56 (3-6-1). A caution on causation: lining up with more defenders correlates with conceding more — '
     'not because deep blocks fail, but because teams already losing tend to pack the defense, so the shape is '
     'partly an effect of the scoreline, not just a cause.'),
    ('Insight 6 — The Counter-Attack Premium. ',
     'Shots taken right after winning the ball back convert at 13.4% versus 10.5% in sustained build-up — a '
     '1.3× premium. Tempo itself, by contrast, is a weak differentiator: goal sequences are barely faster than '
     'non-scoring ones, and the mean and median even disagree — so it is winning possession high up the pitch, '
     'not raw speed, that creates the edge.'),
]
for lead, body in insights:
    add_bullet(body, bold_lead=lead)

add_image('12_formation_boxplot.png', width=4.9,
          caption='Chart 12 — Goals conceded by formation, sorted from most solid (left) to most vulnerable (right).')

# ============ TASK 4: REFLECTION ============
add_heading('4. Reflection', size=15)

add_heading('AI Tool Used', size=12, color=DARKBLUE, space_before=4)
add_body(
    'The analysis was performed with Claude (Anthropic) operating as an agentic assistant. Claude designed '
    'the analysis framework, generated the Python code for data joining, statistics and visualization, '
    'and helped interpret the results in plain language.'
)

add_heading('Advantages', size=12, color=DARKBLUE, space_before=4)
add_bullet('it translated open research questions (“what patterns lead to goals?”) into concrete, executable analysis steps.', bold_lead='Conversational iteration — ')
add_bullet('working pandas, matplotlib and seaborn code that handled a 2.3 GB feature file via column selection.', bold_lead='High-quality code — ')
add_bullet('it understood specialist concepts (the SPADL action format, VAEP goal-sequence labels, conversion rates) and applied them correctly.', bold_lead='Domain knowledge — ')

add_heading('Limitations', size=12, color=DARKBLUE, space_before=4)
add_bullet('Claude cannot see rendered charts, so the human must visually confirm each visualization is correct.', bold_lead='No visual feedback — ')
add_bullet('the coordinate convention (which goal a team attacks) had to be checked manually; an early version measured distance to the wrong goal.', bold_lead='Data assumptions — ')
add_bullet('the dataset records only on-ball events, so off-ball runs and defensive positioning are invisible to this analysis.', bold_lead='Scope of data — ')

add_heading('How AI-Generated Results Were Verified', size=12, color=DARKBLUE, space_before=4)
add_bullet('Key totals (5,104 goals, 11.1% conversion, 2.63 goals/match) were recomputed directly from the raw data and matched.', bold_lead='Cross-checking — ')
add_bullet('each headline statistic was recomputed via different filtering paths (zone counts vs. distance bins) and the results agreed.', bold_lead='Reproducibility — ')
add_bullet('the popular “speed kills” claim was tested and the data did not back it — tempo barely separates scoring from non-scoring sequences, and the mean and median even disagree — so we reported the honest null result rather than the assumed one.', bold_lead='Falsification — ')
add_bullet('results were checked against football intuition (close-range shots convert more, headers are rarer but efficient).', bold_lead='Sanity checks — ')

# Closing line
doc.add_paragraph().paragraph_format.space_after = Pt(2)
close = doc.add_paragraph()
close.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = close.add_run('Full reproducible analysis and all 13 visualizations are provided in the accompanying notebook: '
                   'soccer_goal_analysis.ipynb')
cr.italic = True
cr.font.size = Pt(9)
cr.font.color.rgb = GRAY

# Page footer with page numbers
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Manhaji  •  The Anatomy of a Goal  •  Page ')
run.font.size = Pt(8)
run.font.color.rgb = GRAY
# page number field
fldSimple = OxmlElement('w:fldSimple')
fldSimple.set(qn('w:instr'), 'PAGE')
fp._p.append(fldSimple)

doc.save(OUT)
print(f"Report saved to: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}")
